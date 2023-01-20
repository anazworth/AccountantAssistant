from docx import Document
from docx.text.paragraph import Paragraph
import csv

with open('letter.csv', 'r') as f:
    reader = csv.reader(f)
    your_list = list(reader)

print(your_list)

for person in your_list:

    document = Document('LetterSamples.docx')
    
    document.paragraphs[5].text = person[0]
    document.paragraphs[5].runs[0].font.name = 'Times New Roman'

    document.paragraphs[6].text = person[1]
    document.paragraphs[6].runs[0].font.name = 'Times New Roman'
    
    document.paragraphs[7].text = person[2]
    document.paragraphs[7].runs[0].font.name = 'Times New Roman'
    
    document.paragraphs[11].text = document.paragraphs[11].text.replace('Mickey Mouse', person[0])
    document.paragraphs[11].text = document.paragraphs[11].text.replace('12,000', person[3])
    document.paragraphs[11].runs[0].font.name = 'Times New Roman'
    
    nameToSave = person[0].replace(' ', '')

    document.save(nameToSave + '.docx')
